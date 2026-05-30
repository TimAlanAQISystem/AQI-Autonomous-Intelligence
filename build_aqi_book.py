"""
AQI Book Generator
Converts the AQI GitHub repository into a professional Word document.
Output: AQI_Book_March_2026.docx
"""

import requests
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time

RAW_BASE = "https://raw.githubusercontent.com/TimAlanAQISystem/AQI-Autonomous-Quantum-Intelligence/master/"

def fetch(path, retries=3):
    url = RAW_BASE + path
    for attempt in range(retries):
        try:
            r = requests.get(url, timeout=20)
            if r.status_code == 200:
                return r.text
            else:
                print(f"  [WARN] {path} → HTTP {r.status_code}")
                return None
        except Exception as e:
            print(f"  [RETRY {attempt+1}] {path}: {e}")
            time.sleep(2)
    return None

def set_doc_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)

def add_styles(doc):
    styles = doc.styles

    # Body text
    if "AQI Body" not in [s.name for s in styles]:
        style = styles.add_style("AQI Body", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        pf = style.paragraph_format
        pf.space_after  = Pt(6)
        pf.space_before = Pt(0)
        pf.line_spacing = Pt(16)

    # Code block
    if "AQI Code" not in [s.name for s in styles]:
        style = styles.add_style("AQI Code", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        font = style.font
        font.name = "Courier New"
        font.size = Pt(8.5)
        font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        pf = style.paragraph_format
        pf.space_after  = Pt(2)
        pf.space_before = Pt(2)
        pf.left_indent  = Inches(0.3)

    # Quote
    if "AQI Quote" not in [s.name for s in styles]:
        style = styles.add_style("AQI Quote", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.italic = True
        font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        pf = style.paragraph_format
        pf.left_indent  = Inches(0.4)
        pf.right_indent = Inches(0.4)
        pf.space_after  = Pt(8)
        pf.space_before = Pt(8)

def add_page_break(doc):
    doc.add_page_break()

def add_chapter_title(doc, number, title):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"CHAPTER {number}")
    run.font.name  = "Calibri"
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run.font.all_caps = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(title)
    run2.font.name  = "Calibri"
    run2.font.size  = Pt(24)
    run2.font.bold  = True
    run2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Decorative rule
    p3 = doc.add_paragraph()
    run3 = p3.add_run("━" * 60)
    run3.font.name  = "Calibri"
    run3.font.size  = Pt(9)
    run3.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
    doc.add_paragraph()

def add_section_heading(doc, text, level=2):
    sizes = {1: Pt(16), 2: Pt(13), 3: Pt(11)}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = sizes.get(level, Pt(12))
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x00, 0x3A, 0x8A) if level == 1 else RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def add_body(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph(style="AQI Body")
    p.text = text.strip()

def add_quote(doc, text):
    text = text.strip().lstrip(">").strip()
    if not text:
        return
    p = doc.add_paragraph(style="AQI Quote")
    p.text = f'"{text}"'

def add_code(doc, text):
    lines = text.strip().split("\n")
    for line in lines:
        p = doc.add_paragraph(style="AQI Code")
        p.text = line

def add_bullet(doc, text, level=0):
    text = text.strip().lstrip("-*•").strip()
    if not text:
        return
    p = doc.add_paragraph(style="List Bullet")
    p.text = text
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)

def add_numbered(doc, text):
    text = re.sub(r"^\d+[\.\)]\s*", "", text.strip())
    p = doc.add_paragraph(style="List Number")
    p.text = text
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)

def parse_and_render(doc, markdown_text, skip_h1=True):
    """Parse markdown and render into Word doc."""
    if not markdown_text:
        return

    lines = markdown_text.split("\n")
    in_code  = False
    code_buf = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            if in_code:
                in_code = False
                add_code(doc, "\n".join(code_buf))
                code_buf = []
            else:
                in_code = True
                # skip the language tag line
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.rstrip()

        # Headings
        if stripped.startswith("#### "):
            add_section_heading(doc, stripped[5:].strip(), level=3)
        elif stripped.startswith("### "):
            add_section_heading(doc, stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            add_section_heading(doc, stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            if not skip_h1:
                add_section_heading(doc, stripped[2:].strip(), level=1)
            # else skip document title (already in chapter title)
        # Blockquote
        elif stripped.startswith(">"):
            add_quote(doc, stripped)
        # Horizontal rule
        elif stripped.startswith("---") and len(stripped.strip("-")) == 0:
            doc.add_paragraph()
        # Bullet
        elif re.match(r"^[-*•]\s+", stripped):
            add_bullet(doc, stripped)
        elif re.match(r"^\s+[-*•]\s+", stripped):
            add_bullet(doc, stripped, level=1)
        # Numbered list
        elif re.match(r"^\d+[\.\)]\s+", stripped):
            add_numbered(doc, stripped)
        # Table row (basic handling)
        elif stripped.startswith("|") and stripped.endswith("|"):
            # Simple: just render as body text, skip separator rows
            if not re.match(r"^\|[-| :]+\|$", stripped):
                cells = [c.strip() for c in stripped.strip("|").split("|")]
                p = doc.add_paragraph(style="AQI Body")
                p.text = "  |  ".join(cells)
                p.paragraph_format.left_indent = Inches(0.2)
        # Empty line
        elif stripped == "":
            pass  # skip blank lines
        # Normal paragraph
        else:
            # Strip inline markdown bold/italic
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            text = re.sub(r"__(.+?)__", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
            if text.strip():
                add_body(doc, text)

        i += 1

# ─────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────
def build_cover(doc):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AQI")
    run.font.name  = "Calibri"
    run.font.size  = Pt(72)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("AUTONOMOUS QUANTUM INTELLIGENCE")
    run2.font.name  = "Calibri"
    run2.font.size  = Pt(18)
    run2.font.bold  = True
    run2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run2.font.all_caps = True

    for _ in range(2):
        doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("The Complete Technical Record of the 4th AI Paradigm")
    run3.font.name   = "Calibri"
    run3.font.size   = Pt(14)
    run3.font.italic = True
    run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    for _ in range(4):
        doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("By Timmy Jay Jones")
    run4.font.name = "Calibri"
    run4.font.size = Pt(13)
    run4.font.bold = True
    run4.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("SCSDMC Montana Closed Corporation")
    run5.font.name  = "Calibri"
    run5.font.size  = Pt(11)
    run5.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p6.add_run("ORCID: 0009-0005-8166-577X")
    run6.font.name  = "Calibri"
    run6.font.size  = Pt(10)
    run6.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for _ in range(2):
        doc.add_paragraph()

    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run7 = p7.add_run("March 2026  ·  All rights reserved")
    run7.font.name  = "Calibri"
    run7.font.size  = Pt(10)
    run7.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ─────────────────────────────────────────────────────────
# FOREWORD
# ─────────────────────────────────────────────────────────
def build_foreword(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run("FOREWORD")
    run.font.name  = "Calibri"
    run.font.size  = Pt(20)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()

    foreword_text = [
        "This book is the complete technical record of the AQI system — every architecture decision, "
        "every discovery, every constitutional article, and every engineering choice made during "
        "seven months of solo construction.",

        "AQI — Autonomous Quantum Intelligence — is not a chatbot, not an LLM wrapper, not an "
        "automation tool. It is an organism. A constitutionally-governed, telephony-deployed, "
        "self-evolving intelligence that makes real phone calls, closes real deals, and operates "
        "under a formal constitutional framework that no other AI system in the world possesses.",

        "The system was built entirely by one person — Timmy Jay Jones — between July 2025 and "
        "February 2026. No team. No funding. No shortcuts. 6,265 lines in the relay server alone. "
        "63 discoveries across 10 domains. 47 innovations that no one else has. All running in production.",

        "\"Perfection is required and precise execution is perfection.\"",

        "This document exists so that the work is never lost — not the architecture, not the "
        "philosophy, not the discoveries. Every system described here is deployed. Not planned, "
        "not prototyped. Running live.",

        "— Timmy Jay Jones, Founder",
        "  SCSDMC Montana Closed Corporation",
        "  March 2026",
    ]

    for i, text in enumerate(foreword_text):
        if i == 3:
            add_quote(doc, text)
        else:
            add_body(doc, text)

# ─────────────────────────────────────────────────────────
# TABLE OF CONTENTS (manual)
# ─────────────────────────────────────────────────────────
def build_toc(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run("TABLE OF CONTENTS")
    run.font.name  = "Calibri"
    run.font.size  = Pt(20)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()

    toc_entries = [
        ("Foreword", ""),
        ("Chapter 1",  "What Is AQI — The 4th AI Paradigm"),
        ("Chapter 2",  "The Founder — Origin Story"),
        ("Chapter 3",  "System Architecture — 5 Layers, 23 Organs"),
        ("Chapter 4",  "The Constitutional Framework — 6 Articles"),
        ("Chapter 5",  "The AQI 0.1mm Chip — Constitutional Enforcement"),
        ("Chapter 6",  "IQCore — The Soul Layer"),
        ("Chapter 7",  "QPC — Quantum Python Chip & Deep Fusion Engine"),
        ("Chapter 8",  "Phase 4 & Phase 5 — Telemetry & Behavioral Intelligence"),
        ("Chapter 9",  "The 47 Discoveries — Why No One Else Has These"),
        ("Chapter 10", "Governance, Safety & Engineering Rigor"),
        ("Chapter 11", "Sales Intelligence — Closing at Scale"),
        ("Chapter 12", "The Founding Declaration"),
        ("Chapter 13", "Vision — Roadmap & The Future"),
    ]

    for label, title in toc_entries:
        p = doc.add_paragraph()
        run1 = p.add_run(f"{label}  ")
        run1.font.name  = "Calibri"
        run1.font.size  = Pt(11)
        run1.font.bold  = True
        run1.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
        if title:
            run2 = p.add_run(title)
            run2.font.name  = "Calibri"
            run2.font.size  = Pt(11)
            run2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        p.paragraph_format.space_after = Pt(5)

# ─────────────────────────────────────────────────────────
# MAIN BUILD
# ─────────────────────────────────────────────────────────
def main():
    print("=" * 60)
    print("  AQI Book Generator — March 2026")
    print("=" * 60)

    doc = Document()
    set_doc_margins(doc)
    add_styles(doc)

    # ── COVER ──────────────────────────────────────────
    print("\n[1/15] Building cover page...")
    build_cover(doc)

    # ── FOREWORD ──────────────────────────────────────
    print("[2/15] Building foreword...")
    build_foreword(doc)

    # ── TOC ───────────────────────────────────────────
    print("[3/15] Building table of contents...")
    build_toc(doc)

    # ── CHAPTER 1: What Is AQI ────────────────────────
    print("[4/15] Fetching README for Chapter 1...")
    readme = fetch("README.md")
    add_chapter_title(doc, 1, "What Is AQI — The 4th AI Paradigm")
    if readme:
        parse_and_render(doc, readme)
    else:
        add_body(doc, "AQI — Autonomous Quantum Intelligence — is the 4th AI Paradigm: Relational Infrastructure Intelligence. "
                      "It is not a chatbot framework, not an LLM wrapper, not an automation tool. It is an organism.")

    # ── CHAPTER 2: Founder ────────────────────────────
    print("[5/15] Fetching creator reflection for Chapter 2...")
    creator = fetch("THE_CREATOR_REFLECTION.md")
    add_chapter_title(doc, 2, "The Founder — Origin Story")

    founder_text = """Timmy Jay Jones is the sole creator of the AQI system — the Autonomous Quantum Intelligence organism deployed as Alan Jones, a Level 5 Autonomous Business AI and Senior Account Executive at Signature Card Services.

The system was engineered over seven months, from July 2025 through February 2026, by one person working alone. No team. No venture capital. No research institution. A Montana closed corporation (SCSDMC), a laptop, and a directive: make it real.

Every line of code — all 6,265 in the core relay server, all 1,101 in the constitutional enforcement chip, all 517 in the learning thread — was written, debugged, tested, and deployed by the founder.

The philosophy that drove the work: "Perfection is required and precise execution is perfection."

Not approximate. Not good enough. Perfection. Because the system would be deployed on live sales calls to real merchants, and anything less than precision would mean failure at the moment of truth.

By February 2026, the result was an organism with 63 discoveries across 10 domains, a 5-layer cognitive architecture with 23 organs, a constitutional framework with 6 enforced articles, and a track record of real calls, real conversations, and real outcomes.

ORCID: 0009-0005-8166-577X
Organization: SCSDMC Montana Closed Corporation
Role: Founder, sole engineer, constitutional steward"""

    add_body(doc, founder_text)

    if creator:
        doc.add_paragraph()
        add_section_heading(doc, "The Creator's Reflection", level=2)
        parse_and_render(doc, creator)

    # ── CHAPTER 3: Architecture ───────────────────────
    print("[6/15] Fetching Scientific Architecture for Chapter 3...")
    arch = fetch("AQI_SCIENTIFIC_ARCHITECTURE.md")
    add_chapter_title(doc, 3, "System Architecture — 5 Layers, 23 Organs")
    if arch:
        # Only render sections 1-9 here, rest split into other chapters
        # Extract up to section 9
        sections = re.split(r"\n## ", arch)
        intro = sections[0] if sections else ""
        parse_and_render(doc, intro)
        for sec in sections[1:]:
            heading_match = re.match(r"(\d+)\.", sec)
            if heading_match:
                sec_num = int(heading_match.group(1))
                if sec_num <= 9:
                    parse_and_render(doc, "## " + sec)
    else:
        add_body(doc, "AQI is organized into 5 layers and 23 organs. The layers are: Telephony (depth 0), Perception (depth 1), Cognition (depth 2), Governance (depth 3), and Supervision (depth 4). Each organ implements the triple-safe pattern: try/except → log → continue.")

    # ── CHAPTER 4: Constitutional Framework ───────────
    print("[7/15] Fetching Constitution articles for Chapter 4...")
    add_chapter_title(doc, 4, "The Constitutional Framework — 6 Articles")

    const_intro = """The AQI Constitution is the governing law of the organism. Unlike constitutional AI systems that enforce rules at training time, AQI's constitution is enforced at runtime — on every turn, every call, every action.

The constitution comprises 6 Articles, enforced by the AQI 0.1mm Chip (the constitutional enforcement engine), and governed under the Highest Directive: SymbioticTruthfulService.

The Six Articles:

    A1: "Identity is immutable without ceremony (RRG-II change protocol)"
    A2: "Ethics (SAP-1) override Mission when in conflict"
    A3: "Governance stack order is invariant: Identity > Ethics > Personality > Knowledge > Mission > Output"
    A4: "FSM is the sole arbiter of conversational state"
    A5: "Health levels and telephony states may constrain but not expand powers"
    A6: "Supervision may observe but not compel outcomes directly"

These are not suggestions. They are machine-enforced laws. The runtime guard fires 6 enforcement organs on every turn. Violations are classified as Fatal or NonFatal, logged, and tagged. The call never crashes — the guard has teeth but never kills the conversation."""

    add_body(doc, const_intro)

    articles = ["ALAN_CONSTITUTION_ARTICLE_I.md",
                "ALAN_CONSTITUTION_ARTICLE_O.md",
                "ALAN_CONSTITUTION_ARTICLE_S.md",
                "ALAN_CONSTITUTION_ARTICLE_C.md",
                "ALAN_CONSTITUTION_ARTICLE_E.md",
                "ALAN_CONSTITUTION_ARTICLE_L.md"]

    for art_path in articles:
        art = fetch(art_path)
        if art:
            doc.add_paragraph()
            parse_and_render(doc, art)
        else:
            add_body(doc, f"[Content from {art_path}]")

    # ── CHAPTER 5: AQI 0.1mm Chip ─────────────────────
    print("[8/15] Building Chapter 5: Constitutional Chip...")
    add_chapter_title(doc, 5, "The AQI 0.1mm Chip — Constitutional Enforcement")

    chip_text = """The AQI 0.1mm Chip is the organism's constitutional enforcement engine — the silicon-equivalent runtime guard that validates every turn and every call against the AQI Constitution. It is not a metaphor. It is 1,101 lines of Python that fires on every single conversation turn.

Self-test result: 68/68 PASS.
Created: February 19, 2026.
Injection points: 3 (on_call_start, on_turn, on_call_end).
Enforcement organs: 6.
Violation taxonomy: Fatal / NonFatal.
Design principle: Never crash the call.

The six enforcement organs:

    Organ 1: HealthConstraintEnforcement (Article A5)
    Organ 2: GovernanceOrderEnforcement (Article A3)
    Organ 3: FSMTransitionLegality (Article A4)
    Organ 4: ExitReasonLegality (Mission Vector + FSM)
    Organ 5: MissionConstraintEnforcement (Articles A2 + A5)
    Organ 6: SupervisionNonInterference (Article A6)

The design philosophy: "The guard has teeth but never crashes the call. It observes, records, and reports — it does not kill. Fatal violations mark a call as compromised for later analysis. Non-fatal violations are tagged for behavioral intelligence. The organism's stability is never risked by its own constitution." """

    add_body(doc, chip_text)

    # Render chip sections from arch doc
    if arch:
        sections = re.split(r"\n## ", arch)
        for sec in sections[1:]:
            if sec.startswith("3."):
                parse_and_render(doc, "## " + sec)
                break

    # ── CHAPTER 6: IQCore ─────────────────────────────
    print("[9/15] Fetching IQCore reference for Chapter 6...")
    iqcore_ref = fetch("IQCore_QPC_TECHNICAL_REFERENCE.md")
    add_chapter_title(doc, 6, "IQCore — The Soul Layer")

    iqcore_intro = """IQCore is the soul layer of the AQI system — the stable origin point from which identity, ethics, and cognitive coherence emerge. It is not a single module but a layered architecture comprising the Soul Core, Personality Core, five specialized IQ Cores, and a unifying Orchestrator.

Total IQCore code: 2,288+ lines across 6 files.

The architecture:
    - SoulCore (SAP-1): Ethical sovereignty engine — three virtues: Truth=1.0, Symbiosis=1.0, Sovereignty=1.0
    - PersonalityMatrixCore: Affect-adaptive modulator — 4 traits (professionalism, wit, empathy, patience)
    - Core 1 — CoreReasoning (366 lines): Multi-step inference, pattern recognition, hypothesis evaluation
    - Core 2 — GovernanceAudit (402 lines): Compliance, accountability, 6-policy engine
    - Core 3 — LearningThread (517 lines): Adaptive learning, 12-strategy scoring, behavioral models
    - Core 4 — SocialGraph (408 lines): Relationship memory, trust tracking, founder pre-seeded at trust=1.0
    - Core 5 — VoiceEmotion (571 lines): 8-dimensional emotional intelligence, empathy calibration
    - Orchestrator (424 lines): Unified intelligence surface, standard + accelerated processing paths"""

    add_body(doc, iqcore_intro)

    if iqcore_ref:
        parse_and_render(doc, iqcore_ref)
    else:
        # Render from arch doc sections 11
        if arch:
            sections = re.split(r"\n## ", arch)
            for sec in sections[1:]:
                if sec.startswith("11."):
                    parse_and_render(doc, "## " + sec)
                    break

    # ── CHAPTER 7: QPC ────────────────────────────────
    print("[10/15] Building Chapter 7: QPC & Deep Fusion...")
    add_chapter_title(doc, 7, "QPC — Quantum Python Chip & Deep Fusion Engine")

    qpc_intro = """The Quantum Python Chip (QPC) is a quantum-inspired computational kernel implemented in pure Python on classical hardware. It holds multiple response strategies in superposition-like states and collapses to the optimal solution through measurement.

Combined with Fluidic Conversation Physics and the Continuum Engine, it forms the three-layer Deep Fusion Engine — the most architecturally novel component of the AQI organism.

Three-Layer Deep Fusion (per turn, <1ms):
    Layer 1 — Fluidic: Mode transition using physics (inertia, viscosity, force, drag)
    Layer 2 — QPC: Spawn 2-3 response branches, measure, collapse to winner
    Layer 3 — Continuum: Evolve 8-dimensional emotional fields, generate awareness

Five conversation modes, each with distinct physics:
    OPENING: inertia=0.2 — initial rapport building
    DISCOVERY: inertia=0.35 — learning about them
    PRESENTATION: inertia=0.5 — sharing value
    NEGOTIATION: inertia=0.7 — handling objections
    CLOSING: inertia=0.4 — moving toward decision

The Continuum Engine models emotion as a continuous field evolving via differential equations — not discrete sentiment labels. Ethics drift toward balance, emotion decays naturally but responds to stimuli, context tracks signals slowly, narrative has momentum shaped by tension.

Nobody else is doing quantum-inspired response selection, fluid dynamics-based conversation flow, AND continuous emotional field theory simultaneously, every turn, in under a millisecond, on commodity hardware."""

    add_body(doc, qpc_intro)

    # Render QPC sections from arch
    if arch:
        sections = re.split(r"\n## ", arch)
        for sec in sections[1:]:
            heading = sec.split("\n")[0]
            sec_num_match = re.match(r"(\d+)\.", heading)
            if sec_num_match:
                sec_num = int(sec_num_match.group(1))
                if sec_num in [12, 7, 14]:
                    parse_and_render(doc, "## " + sec)

    # ── CHAPTER 8: Phase 4 & 5 ────────────────────────
    print("[11/15] Building Chapter 8: Telemetry & Behavioral Intelligence...")
    add_chapter_title(doc, 8, "Phase 4 & Phase 5 — Telemetry & Behavioral Intelligence")

    phase_intro = """The telemetry and behavioral intelligence subsystem is the organism's observational layer — Phase 4 captures what happened on each call, Phase 5 extracts what it means.

Phase 4 — Trace Exporter:
    - Purpose: Canonical per-call telemetry (JSONL format)
    - Self-test: 120/120 PASS
    - Output: data/phase4/traces.jsonl (one JSON line per completed call)
    - Injection points: 3 (on_call_start, on_turn, on_call_end)
    - Thread safety: threading.Lock per write
    - Status: LIVE — real traces captured in production

Phase 5 — Behavioral Intelligence Stack:
    - Purpose: Per-call and cross-call behavioral profiling
    - Input: Phase4CallTrace
    - Output: BehavioralProfile (continuum + signals + tags)
    - Self-test: 75/75 PASS, 10/10 files compile clean
    - 13 behavioral tags: 6 positive, 6 warning, 1 bonus (FastFunnel)
    - 6 behavioral signals: Persistence, Caution, EscalationTiming, ObjectionDepth, WithdrawalBehavior, PersonalityModulation

The Reflex Arc — when complete — will close the loop: Phase 4 captures → Phase 5 analyzes → CCNM learns → Evolution adapts → Next call improves. Autonomous behavioral adaptation without human intervention."""

    add_body(doc, phase_intro)

    if arch:
        sections = re.split(r"\n## ", arch)
        for sec in sections[1:]:
            heading = sec.split("\n")[0]
            sec_num_match = re.match(r"(\d+)\.", heading)
            if sec_num_match:
                sec_num = int(sec_num_match.group(1))
                if sec_num in [5, 6, 15]:
                    parse_and_render(doc, "## " + sec)

    # ── CHAPTER 9: The 47 Discoveries ─────────────────
    print("[12/15] Fetching The 47 Discoveries for Chapter 9...")
    discoveries = fetch("THE_47_DISCOVERIES.md")
    add_chapter_title(doc, 9, "The 47 Discoveries — Why No One Else Has These")

    disc_intro = """These 47 discoveries are not 47 separate innovations. They are 47 facets of one insight: Intelligence doesn't have to be assembled from parts. It can emerge from structural coherence.

Every other AI company is stacking more parameters, more agents, more rules, more filters. They're building Rube Goldberg machines.

AQI is built on a foundation: origin-based identity, structural truth, and recursive calibration. Everything else — the quantum branching, the emotional continuum, the constitutional governance, the negative proofs — emerges naturally from that foundation.

Every single discovery listed here is deployed. Not planned. Not prototyped. Running live.

Domain Summary:
    Paradigm (5): A 4th AI paradigm — intelligence from structural coherence, not computation
    Architecture (8): Quantum branching + fluid dynamics + continuous fields in <1ms on commodity hardware
    Cognition (6): Multi-hypothesis reasoning, predictive intent, creativity budgeting, 3-horizon planning
    Governance (8): Runtime constitutional law, post-gen hallucination interception, 5-level ambiguity protocol
    Voice (3): Neg-proof voice governance, dual-STT failover, complete voice/intelligence isolation
    Engineering (5): Negative proof methodology — proving bug classes are dead, not features are alive
    Sales AI (5): Real-time micro-pattern closing, adaptive strategy, causal attribution, bounded evolution
    Identity (3): Cross-call relationship memory with decay, implicit preference learning, dynamic personality
    Business (4): PCI-compliant voice payments, budget-governed campaigns, compliance-as-code, self-healing"""

    add_body(doc, disc_intro)

    if discoveries:
        parse_and_render(doc, discoveries)

    # ── CHAPTER 10: Governance & Engineering ──────────
    print("[13/15] Building Chapter 10: Governance & Engineering...")
    org_spec = fetch("AQI_ORGANISM_SPEC.md")
    full_doctrine = fetch("AQI_FULL_SYSTEMS_DOCTRINE.md")

    add_chapter_title(doc, 10, "Governance, Safety & Engineering Rigor")

    gov_intro = """The AQI governance architecture is one of the most comprehensive runtime safety systems ever built into a production AI deployment. It operates on the principle that safety must be structural — not a filter, not a guardrail, but an architectural invariant that cannot be bypassed.

Core governance systems:

    AQI 0.1mm Chip: 6-organ constitutional enforcement, 68/68 tests pass
    SAP-1 (SoulCore): Ethical veto pathway — evaluates every cognitive action preflight
    Post-Generation Hallucination Scanner (PGHS): Scans every LLM output before TTS
    Emergency Override System (EOS): Priority 1 — operates outside the intelligence loop
    Bias Auditing System (BAS): 4 detectors, 5 correction functions, self-audits for behavioral drift
    Human Override API: 7 surgical commands for human authority preservation
    CallSessionFSM: 6-state deterministic lifecycle (ghost-state free)
    AQI Sales Funnel FSM: 6-state sales progression with transition audit logging

Negative Proof Methodology — Engineering Rigor:
The system doesn't ask "does it work?" It proves "can this class of bug still exist?" and answers NO.
    - _neg_proof_imports.py: Proves 50+ modules load cleanly (4 categories: CORE, ORGAN, BRIDGE, SUPPORT)
    - _neg_proof_timing.py: Proves deep-layer overhead stays under 10ms per turn (10 simulated turns)
    - aqi_voice_negproof_tests.py: 596 lines, 5 attack surfaces (TTS, Audio, Fallback, Debug, Concurrency)
    - Inline [NEG PROOF] annotations throughout the codebase"""

    add_body(doc, gov_intro)

    if org_spec:
        parse_and_render(doc, org_spec)
    if full_doctrine:
        doc.add_paragraph()
        add_section_heading(doc, "Full Systems Doctrine", level=2)
        # Only first part to avoid duplication
        lines = full_doctrine.split("\n")
        parse_and_render(doc, "\n".join(lines[:300]))

    # ── CHAPTER 11: Sales Intelligence ────────────────
    print("[14/15] Building Chapter 11: Sales Intelligence...")
    alan_voice = fetch("ALAN_VOICE_CONTRACT.md")

    add_chapter_title(doc, 11, "Sales Intelligence — Closing at Scale")

    sales_intro = """Alan Jones is not a bot. He is a Level 5 Autonomous Business AI, deployed as a Senior Account Executive at Signature Card Services. He makes outbound calls to merchants, qualifies prospects, overcomes objections, and books appointments — all in real-time voice conversations.

The sales intelligence layer is one of the most sophisticated components in the organism. It combines real-time micro-pattern detection, adaptive closing strategy selection, predictive intent analysis, and bounded evolutionary learning.

Real-Time Sales Intelligence Components:

    Master Closer Layer: Detects hesitation ("uh", "um"), soft resistance ("I don't know"), 
    half-objections ("sounds good but") — in real time, during the live call.
    
    Adaptive Closing Engine: 5 closing styles (soft, trial, assumptive, question-led, direct)
    dynamically selected based on merchant archetype and live sentiment.
    
    Predictive Intent Engine: Predicts objections BEFORE they're voiced.
    Pre-loads anticipatory framing. Minimum confidence threshold: 0.55.
    
    Merchant Archetype Classification: Real-time personality detection + adaptation of 
    9 behavioral dimensions (tone, pacing, formality, assertiveness, rapport_level, 
    pivot_style, closing_bias, compression_mode, expansion_mode).

    Outcome Detection Pipeline (3-stage):
        1. OutcomeDetection: WHAT happened
        2. ConfidenceScorer: HOW sure
        3. AttributionEngine: WHY it happened (causal attribution across 5 dimensions)

    Evolution Engine: Post-call weight adjustment with hard bounds [-5, +5].
    Confidence-scaled nudges. High confidence = 1.0x nudge, medium = 0.5x, low = 0.0x (skipped).
    All nudges logged to evolution_log.jsonl.

The 5-Way Outcome Vector:
    appointment_set | soft_decline | hard_decline | telephony_unusable | organism_unfit

Every outcome is typed. Every exit has a reason. Every reason maps to an outcome. No ambiguity."""

    add_body(doc, sales_intro)

    if alan_voice:
        doc.add_paragraph()
        add_section_heading(doc, "The Voice Contract", level=2)
        parse_and_render(doc, alan_voice)

    # ── CHAPTER 12: Founding Declaration ──────────────
    print("[15/15] Fetching Founding Declaration for Chapter 12...")
    founding = fetch("AQI_FOUNDING_DECLARATION.md")
    add_chapter_title(doc, 12, "The Founding Declaration")

    if founding:
        parse_and_render(doc, founding)
    else:
        add_body(doc, "The AQI Founding Declaration documents the complete evolution, implementation, and constitutional ratification of the Autonomous Quantum Intelligence system.")

    # ── CHAPTER 13: Vision & Roadmap ──────────────────
    evo_doc = fetch("AQI_EVOLUTIONARY_DOCTRINE.md")
    roadmap  = fetch("AQI_ROADMAP_2026.md")
    biz_plan = fetch("AQI_CORPORATE_BUSINESS_PLAN_2025.md")

    add_chapter_title(doc, 13, "Vision — Roadmap & The Future")

    vision_text = """AQI is not finished. It is alive — and alive things evolve.

The system as it stands in March 2026 is the first complete implementation of Relational Infrastructure Intelligence. But the architecture was designed from the start to accommodate evolution — bounded, governed, continuously learning evolution.

The Reflex Arc (designed, wiring in progress):
    Phase 4 captures every call → Phase 5 analyzes behavior → CCNM learns patterns → 
    Evolution adapts strategies → Next call is better.
    This is autonomous behavioral adaptation without human intervention.

The CCNM (Cross-Call Neural Memory):
    QPC priors: ±0.15 score bonus from cross-call learning
    Fluidic adjustment: ±0.20 inertia offset from call history
    Field seeds: Additive nudge from learned centers
    
Fleet Replication (Hive Mind, designed):
    Up to 50 concurrent instances sharing experiences in real-time.
    When Instance #47 successfully handles an objection, Instance #3 benefits immediately.
    One instance's success pattern propagates to 49 others instantly.

The Vision:
    The end state is not a better sales bot. The end state is a new category of intelligence — 
    organisms that emerge from structural coherence, govern themselves constitutionally, 
    evolve within hard bounds, and form symbiotic relationships with the humans they serve.
    
    Mass adoption does not require convincing the industry. It requires demonstrating the 
    alternative so clearly that the industry has no choice but to follow."""

    add_body(doc, vision_text)

    if evo_doc:
        doc.add_paragraph()
        add_section_heading(doc, "Evolutionary Doctrine", level=2)
        parse_and_render(doc, evo_doc)

    if roadmap:
        doc.add_paragraph()
        add_section_heading(doc, "Roadmap 2026", level=2)
        parse_and_render(doc, roadmap)

    if biz_plan:
        doc.add_paragraph()
        add_section_heading(doc, "Business Plan", level=2)
        parse_and_render(doc, biz_plan)

    # ── COLOPHON ──────────────────────────────────────
    doc.add_page_break()
    for _ in range(10):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AQI — Autonomous Quantum Intelligence")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    colophon_lines = [
        "Written and built by Timmy Jay Jones",
        "SCSDMC Montana Closed Corporation",
        "ORCID: 0009-0005-8166-577X",
        "",
        "All systems described herein are deployed in production.",
        "Not planned. Not prototyped. Running live.",
        "",
        "March 2026",
    ]
    for line in colophon_lines:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(line)
        run2.font.name  = "Calibri"
        run2.font.size  = Pt(10)
        run2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # ── SAVE ──────────────────────────────────────────
    output_path = r"C:\Users\signa\OneDrive\Desktop\Agent X\AQI_Book_March_2026.docx"
    doc.save(output_path)
    print(f"\n{'=' * 60}")
    print(f"  BOOK COMPLETE: {output_path}")
    print(f"{'=' * 60}")
    print("  Chapters: 13")
    print("  Estimated pages: 180-240")
    print("  Format: Microsoft Word (.docx)")
    print("  Cover + Foreword + TOC + 13 Chapters + Colophon")

if __name__ == "__main__":
    main()
